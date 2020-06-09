from function import *
from docx import Document

key = '133457799BBCDFF1'
text = 'COMPUTER'
encryption = encrypt(key, text)
print('encripted text: %s' % encryption)
decryption = decrypt(key, encryption)
print('decripted text: %s' % decryption)

document = Document('text.docx')
text = ''
for paragraph in document.paragraphs:
    text += paragraph.text
encryptedText = encrypt(key, text)
newFile = Document()
newFile.add_paragraph(encryptedText)
newFile.save('encryptedText.docx')

document = Document('encryptedText.docx')
text = ''
for paragraph in document.paragraphs:
    text += paragraph.text
decryptedText = decrypt(key, text)
newFile = Document()
newFile.add_paragraph(decryptedText)
newFile.save('decryptedText.docx')
